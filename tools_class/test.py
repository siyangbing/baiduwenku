import os
import time

from selenium import webdriver
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from scrapy import Selector
import json
import requests
from my_fake_useragent import UserAgent
import docx
from docx.shared import Inches
import cv2
from pptx import Presentation
from pptx.util import Inches
from tools_class.xpath_string import click_xpath_dict, file_attribute_xpath, content_dict, next_page_cancel_dict

url = "https://wenku.baidu.com/view/369c753aa65177232f60ddccda38376bae1fe05b.html?fr=search-4-X-income1&fixfr=zT1apvO%2BF0K%2B1ichceLdSw%3D%3D"
# url = "https://wenku.baidu.com/view/3de365cc6aec0975f46527d3240c844769eaa0aa.html?fr=search"


# url = "https://wenku.baidu.com/view/6e70d05abe23482fb4da4cad.html"

# 下载图片类
class DownloadImg():
    def __init__(self):
        self.ua = UserAgent()

    def download_one_img(self, img_url, saved_path):
        # 下载图片
        header = {
            "User-Agent": "{}".format(self.ua.random().strip()),
            'Connection': 'close'}
        r = requests.get(img_url, headers=header, stream=True)
        print("请求图片状态码　{}".format(r.status_code))  # 返回状态码
        if r.status_code == 200:  # 写入图片
            with open(saved_path, mode="wb") as f:
                f.write(r.content)
            print("download {} success!".format(saved_path))
        del r
        return saved_path


# 浏览器操纵类
class ControlChrome():
    # 初始化并打开浏览器
    def __init__(self, chromedriver_path="./chromedriver"):
        mobile_emulation = {"deviceName": "Galaxy S5"}
        capabilities = DesiredCapabilities.CHROME
        capabilities['loggingPrefs'] = {'browser': 'ALL'}
        options = webdriver.ChromeOptions()
        options.add_experimental_option("mobileEmulation", mobile_emulation)
        self.brower = webdriver.Chrome(executable_path=chromedriver_path, desired_capabilities=capabilities,
                                       options=options)

    def go_url(self, url):
        # 启动浏览器，打开需要下载的文档网页
        self.brower.get(url)

    def scrolled_to_end(self):
        # 滑动滚轮到浏览器页面最底部
        js = "return action=document.body.scrollHeight"
        # 初始化现在滚动条所在高度为0
        height = 0
        # 当前窗口总高度
        new_height = self.brower.execute_script(js)

        while height < new_height:
            # 将滚动条调整至页面底部
            for i in range(height, new_height, 200):
                self.brower.execute_script('window.scrollTo(0, {})'.format(i))
                time.sleep(0.5)
                print("滑动滚轮！")
            height = new_height
            new_height = self.brower.execute_script(js)


        # self.brower.execute_script("window.scrollTo(0,document.body.scrollHeight)")

    def click_ele(self, click_xpath):
        # 单击指定控件
        click_ele = self.brower.find_elements_by_xpath(click_xpath)
        if click_ele:
            click_ele[0].location_once_scrolled_into_view  # 滚动到控件位置
            self.brower.execute_script('arguments[0].click()', click_ele[0])  # 单击控件，即使控件被遮挡，同样可以单击
        else:
            print("not found {}".format(click_ele))
        return click_ele

    def get_html(self):
        sel = Selector(text=self.brower.page_source)
        return sel


class FilterXpath():
    def __init__(self, click_xpath_dict, file_attribute_xpath, content_dict, next_page_cancel_dict):
        self.click_xpath_dict = click_xpath_dict
        self.file_attribute_xpath = file_attribute_xpath
        self.content_dict = content_dict
        self.next_page_dict_cancel = next_page_cancel_dict

    def get_filetype(self, sel):
        # 获取文件类型
        file_type_xpath = self.file_attribute_xpath["file_type"][0]
        file_type = sel.xpath(file_type_xpath).extract()[0]
        file_type = file_type.split(' ')[-1]
        return file_type

    def get_filename(self, sel):
        # 获取文件名称
        file_name_xpath = self.file_attribute_xpath["file_name"][0]
        file_name = sel.xpath(file_name_xpath).extract()[0]
        file_name = file_name.split('-')[-2]
        return file_name

    def get_all_page(self, cc):
        # 循环点击浏览，展示所有页面
        load_more_page = self.click_xpath_dict["load_more_page"][0]
        load_next_page = self.click_xpath_dict["load_more_page"][1]
        end_page = self.click_xpath_dict["end_page"][0]

        cc.click_ele(load_more_page)
        login_cancel_xpath = self.click_xpath_dict["login_cancel_xpath"][0]
        cc.click_ele(login_cancel_xpath)

        while True:
            cc.click_ele(load_next_page)
            cc.scrolled_to_end()
            print("点击加载更多！")
            b = cc.brower.find_elements_by_xpath(end_page)
            if b != []:
                break

    def get_content(self,sel):
        content_xpath = self.content_dict["doc"][0]
        file_content= sel.xpath(content_xpath).extract()
        return file_content


class CreateFile():
    def __init__(self, filetype, file_title, file_content):
        self.filetype = filetype
        self.file_title = file_title
        self.file_content = file_content

    def create_doc(self,title,contents):
        document = docx.Document()  # 创建word文档
        document.add_heading(title, 0)  # 添加标题
        content_all = ""
        for content_one in contents:
            if content_one==" ":
                content_all += ('\n'+content_one)
            else:
                content_all += content_one
            # content_txt_one = '*'.join(content_one.xpath(xpath_content_one).extract())
        document.add_paragraph(content_all)
        document.save(os.path.join("../doc/", '{}.docx'.format(title)))


if __name__ == "__main__":
    fx = FilterXpath(click_xpath_dict, file_attribute_xpath, content_dict, next_page_cancel_dict)
    cc = ControlChrome()
    cc.go_url(url)
    # cc.scrolled_to_end()
    time.sleep(3)
    # 获取文档的所有页面
    fx.get_all_page(cc)
    # cc.scrolled_to_end()
    # 获取当前html
    sel = cc.get_html()
    # 获取文件类型
    file_type_str = fx.get_filetype(sel)
    # 获取文件名称
    filename_str = fx.get_filename(sel)

    file_content = fx.get_content(sel)

    cf = CreateFile(file_type_str, filename_str, file_content)
    cf.create_doc(filename_str,file_content)
    print("")
